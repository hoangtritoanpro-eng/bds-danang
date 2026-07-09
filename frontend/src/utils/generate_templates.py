import os
from docx import Document
import json

def get_text(doc_path):
    doc = Document(doc_path)
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace('\n', ' '))
            text.append(" | ".join(row_data))
    return "\n".join(text)

muaban = get_text(r"C:\Users\hoang\Downloads\TOAN\duanvinhbds\hosohopdongmuaban\hop dong mua ban mau.docx")
thue = get_text(r"C:\Users\hoang\Downloads\TOAN\duanvinhbds\hosohopdongmuaban\hop dong thue mau.docx")

js_content = f"""// Generated contract templates
export const MUA_BAN_TEMPLATE = `{muaban.replace('`', '\\`')}`;

export const THUE_KHO_TEMPLATE = `{thue.replace('`', '\\`')}`;
"""

out_path = r"C:\Users\hoang\Downloads\TOAN\duanvinhbds\bds-danang\frontend\src\utils\contractTemplates.js"
with open(out_path, "w", encoding="utf-8") as f:
    f.write(js_content)

print(f"Successfully wrote {out_path}")
